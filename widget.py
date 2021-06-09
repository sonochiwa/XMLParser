import csv
import docx
from item import Ui_Item
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from form import Ui_Form
from db.base import session
from db.models import RKN
from db.utils import get_field_type
from datetime import datetime
from XMLparser import parser
from sqlalchemy.sql.sqltypes import String, Integer, Boolean


def select(**kwargs):
    with session() as s:
        query_set = s.query(RKN).filter_by(**kwargs)        
    return query_set

def convert(sqltype, value):
    if value is None:
        return value
    if isinstance(sqltype, String):
        return value
    elif isinstance(sqltype, Integer):
        return int(value)
    elif isinstance(sqltype, Boolean):
        return bool(int(value))

class Item(QWidget, Ui_Item):
    def __init__(self):
        super(Item, self).__init__()
        self.setupUi(self)
        self.comboBox.addItems(RKN.__table__.columns.keys()[1:])

class Main(QWidget, Ui_Form):
    def __init__(self):
        super(Main, self).__init__()
        self.setupUi(self)
        self.file_path = ''
        self.pushButton_1.clicked.connect(self.get_xml_file)
        self.pushButton_2.clicked.connect(self.parse_XML)
        self.pushButton_3.clicked.connect(self.add_item)
        self.pushButton_4.clicked.connect(self.btn_select)
        self.pushButton.clicked.connect(self.del_item)
        self.contentLayout = QVBoxLayout()
        self.areaContent.setLayout(self.contentLayout)
        self.setUpContent()

    def del_item(self):
        self.contentLayout.itemAt(self.contentLayout.count()-1).widget().deleteLater()

    def setUpContent(self):
        if not self.contentLayout.count():
            self.add_item()

    def add_item(self):
        w = Item()
        self.contentLayout.addWidget(w)

    def btn_select(self):
        data = {}
        for elem in self.areaContent.children():
            if isinstance(elem, QWidget):
                key = elem.comboBox.currentText()
                value = elem.lineEdit_2.text()    
                sqltype = get_field_type(RKN, key)
                value = convert(sqltype, value)
                data[key] = value
        result = select(**data)
        self.info('Найдено записей: {}'.format(result.count()))
        with open('dump.csv', 'w', newline='') as f:            
            writer = csv.writer(f)
            writer.writerow([column.name for column in RKN.__mapper__.columns])
            [writer.writerow([getattr(curr, column.name) for column in RKN.__mapper__.columns]) for curr in result]

        mydoc = docx.Document()
        mydoc.add_paragraph('Дата - {}'.format(datetime.now().strftime('%d.%m.%Y')))
        mydoc.add_paragraph('Время - {}'.format(datetime.now().strftime('%H:%M')))
        mydoc.add_paragraph('Поле - {}, записей - {}'.format(', '.join(data.keys()), result.count()))
        mydoc.save("report.docx")

    def get_xml_file(self):
        _filter = '*.xml'
        self.file_path, _ = QFileDialog.getOpenFileName(self, 'Set folder', 'c:/tmp', _filter)
        filename = self.file_path.split("/")[-1]
        self.lineEdit_1.setText(filename)

    def parse_XML(self):
        if self.file_path:
            self.info('Парсинг в процессе...')
            self.xml_to_db()
            self.info('Готово!')
            self.info('Всего записей прочитано: {}'.format(self.all_read))
            self.info('Всего записей вставлено: {}'.format(self.records))
            self.info('Время обработки данных: {}'.format(self.seconds))

    def info(self, text):
        self.textBrowser.append(text)
        self.textBrowser.repaint()

    def xml_to_db(self):
        rcode = self.spinBox.value()
        RKN.metadata.create_all()
        path = self.file_path
        self.all_read = 0
        self.records = 0
        start = datetime.now()
        with session() as s:
            for record in parser(path):
                self.all_read += 1
                if record['region_code'] != str(rcode): continue
                rkn = RKN()
                for key, value in record.items():
                    if hasattr(rkn, key):
                        sqltype = get_field_type(RKN, key)
                        value = convert(sqltype, value)
                        setattr(rkn, key, value)
                s.add(rkn)
                s.flush()
                self.records += 1
                self.seconds = datetime.now() - start
    
if __name__ == '__main__':
    app = QApplication([])
    w = Main()
    w.show()
    app.exec_()